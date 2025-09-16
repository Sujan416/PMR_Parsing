from docx import Document
import json
import os
from jinja2 import Environment, meta, Template
from docx.enum.text import WD_COLOR_INDEX
import sys


def load_all_json_data(foldejson_data_folder_pathr_path):
    """Load all JSON files into a list of dictionaries"""
    data_list = []
    for filename in os.listdir(json_data_folder_path):
        if filename.endswith('.txt'):
            path = os.path.join(json_data_folder_path, filename)
            with open(path, 'r', encoding='utf-8') as f:
                try:
                    data = json.load(f)
                    data_list.append(data)
                except json.JSONDecodeError:
                    print(f"Invalid JSON: {filename}")
    return data_list



def try_render_with_data(template_str, data_sources):
    """Try rendering the template string using any of the JSON data sources"""
    env = Environment()
    ast = env.parse(template_str)
    variables = meta.find_undeclared_variables(ast)

    for data in data_sources:
        try:
            # Check if all variables exist in this data source
            all_vars_exist = True
            for var in variables:
                if resolve_variable(var, data) is None:
                    all_vars_exist = False
                    break
            if all_vars_exist:
                template = Template(template_str)
                return template.render(data)
        except Exception:
            continue
    # If no match, return original string
    return template_str

def resolve_variable(var_path, data):
    """Safely resolve nested variable path like 'user.name' from a dictionary"""
    try:
        parts = var_path.split('.')
        for part in parts:
            if isinstance(data, dict):
                data = data.get(part)
            else:
                return None
        return data
    except Exception:
        return None

def process_docx(doc, data_sources):
    """Replace Jinja2 templates in all table cells with resolved values"""
    highlight_words = ['f', 'stopped', 'failed', 'error' ]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text.strip()
                if '{' in original_text and '}' in original_text:
                    new_text = try_render_with_data(original_text, data_sources)
                    # cell.text = new_text.strip()
                else:
                    new_text = original_text
                cell.text = ''  # Clear existing content
                p = cell.paragraphs[0]
                run = p.add_run(new_text.strip())

                if new_text.strip().lower() in highlight_words:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW



if __name__ == "__main__":
    if len(sys.argv) > 2:
        template_file_path = sys.argv[1]
        output_file = sys.argv[2]
        json_data_folder_path = sys.argv[3]
        
        json_data_list = load_all_json_data(json_data_folder_path)
        doc = Document(template_file_path)
        process_docx(doc, json_data_list)
        doc.save(output_file)
        print(f"Saved filled document to: {output_file}")

    else:
        print("Usage: python main.py <template_file_path> <output_file> <json_data_folder_path>")
        exit(1)
    # template_file_path = r"C:\\ProjectFiles\\Maintenance\\Projects\\PMR_Parsing\\docs\\SICE-WCXS&M-M4M5-RPT-SEP25-D07_Template.docx"
    # output_file = r"C:\\ProjectFiles\\Maintenance\\Projects\\PMR_Parsing\\docs\\SICE-WCXS&M-M4M5-RPT-SEP25-D07.docx"
    # json_data_folder_path = r"C:\\ProjectFiles\\Maintenance\\Projects\\PMR_Parsing\\data\\2025-09-12_WCX3A"


